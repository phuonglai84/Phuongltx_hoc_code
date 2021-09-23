#!/usr/bin/python -tt
# Copyright 2010 Google Inc.
# Licensed under the Apache License, Version 2.0
# http://www.apache.org/licenses/LICENSE-2.0
# Quang Le - Techmaster.vn - 09/2021
########

"""Wordcount exercise

Hàm main() đã được định nghĩa hoàn chỉnh ở dưới. Bạn phải viết hàm print_words()
và print_top() mà sẽ được gọi từ main().

1. Với đối số --count, viết hàm print_words(filename) đếm số lần xuất hiện của mỗi từ 
trong file đầu vào và in ra theo định dạng sau:
word1 count1
word2 count2
...

In danh sách trên theo thứ tự từ điển các từ (python sẽ sắp xếp dấu câu đứng trước
các chữ cái nên cũng không thành vấn đề). Lưu tất cả các từ dưới dạng chữ thường,
vì vậy 'The' và 'the' được tính là cùng một từ.

2. Với đối số --topcount, viết hàm print_top(filename) tương tự như print_words()
nhưng chỉ in ra 20 từ thông dụng nhất sắp xếp theo từ thông dụng nhất ở trên cùng.

Tùy chọn: định nghĩa một hàm helper để tránh lặp lại code trong các hàm 
print_words() và print_top().

"""

import sys
print("Bài 3: Đếm từ")
# +++your code here+++
def print_words(filename):
    from docx import Document

    document = Document('d:\\Learning\\Python\\Phuongltx_hoc_code\\pltxpython/'+filename)
    allText = []
    for para in document.paragraphs:
        para_text = para.text.lower()
        word = para_text.split()
        allText.extend(word)
        allText2 = sorted(allText, reverse=True)    # sort ngược, để khi chạy for key lấy ra value tạo ra string word_count đúng chiều xuôi

    # print(allText)
    # print(sorted(allText))
    d={}
    for i in allText2:
        allText2.count(i)
        d[i] = allText2.count(i)
    d_top = sorted(d.items(), key = lambda x: x[1], reverse=True )
    # print(d)
    # print(d_top)
    word_count = ""
    for j in d.keys():
        word_count = str(j) + "\t" + str(d[j]) + "\n" + word_count
        # word_count_reverse = word_count[::-1]
    print(word_count)
    return word_count

def print_top(filename):
    from docx import Document

    document = Document('d:\\Learning\\Python\\Phuongltx_hoc_code\\pltxpython/'+filename)
    allText = []
    for para in document.paragraphs:
        para_text = para.text.lower()
        word = para_text.split()
        allText.extend(word)
        allText2 = sorted(allText, reverse=True)    # sort ngược, để khi chạy for key lấy ra value tạo ra string word_count đúng chiều xuôi

    # print(allText)
    # print(sorted(allText))
    d={}
    for i in allText2:
        allText2.count(i)
        d[i] = allText2.count(i)
    d_valuesort = sorted(d.items(), key = lambda x: x[1] )
    # print(d)
    # print(d_valuesort)
    if len(d_valuesort) > 20:
        d_valuetop = d_valuesort[0:20] 
    else:
        d_valuetop =  d_valuesort
    # print(d_valuetop)
    word_top = ""
    for j in d_valuetop:
        word_top = str(j[0]) + "\t" + str(j[1]) + "\n" + word_top
    print(word_top)
    return word_top

###

# This basic command line argument parsing code is provided and
# calls the print_words() and print_top() functions which you must define.
def main():
  if len(sys.argv) != 3:
    print('usage: ./wordcount.py {--count | --topcount} file')
    sys.exit(1)

  option = sys.argv[1]
  filename = sys.argv[2]
  if option == '--count':
    print_words(filename)
  elif option == '--topcount':
    print_top(filename)
  else:
    print('unknown option: ' + option)
    sys.exit(1)

if __name__ == '__main__':
  main()

# Điểm thi học kỳ của sinh viên được lưu ở định dạng 1 tuple có 3 phần tử (m1, m2, e) gồm:
# m1 = midterm1
# m2 = midterm2
# e = endterm
# Cho một list gồm danh sách điểm thi của sinh viên 1 lớp. 
# Viết chương trình Python để sắp xếp danh sách trước theo thứ tự tăng dần theo phần tử cuối cùng trong mỗi tuple 
# (sắp xếp theo điểm cuối kỳ - endterm tăng dần).
# vd sort_list_last([(1, 2, 5), (9, 1, 2), (6, 4, 4), (3, 2, 3), (10, 2, 1)]) == [(10, 2, 1), (9, 1, 2), (3, 2, 3), (6, 4, 4), (1, 2, 5)]
print("Bài 2: Sắp xếp điểm thi")
score_list = [(1, 2, 5), (9, 1, 2), (6, 4, 4), (3, 2, 3), (10, 2, 1)]

score_list_sorted = sorted(score_list)
print(score_list_sorted)

