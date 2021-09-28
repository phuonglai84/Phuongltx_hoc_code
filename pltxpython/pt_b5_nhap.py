# from os import replace
# import docx2txt


# filename = "text.txt"
# file = open("d:\\Learning\\Python\\Phuongltx_hoc_code\\pltxpython/" + filename, 'r')
# content = file.read()
# content_new = content.replace(".","")
# print(content_new)
# d = {}
def print_words(filename):
    from docx import Document

    document = Document('d:\\Learning\\Python\\Phuongltx_hoc_code\\pltxpython/'+filename)
    allText = []
    for para in document.paragraphs:
        para_text = para.text.lower()
        word = para_text.split()
        allText.extend(word)
        allText2 = sorted(allText, reverse=True)    # sort ngược, để khi chạy for key lấy ra value tạo ra string word_count đúng chiều xuôi

    # print(allText)
    # print(sorted(allText))
    d={}
    for i in allText2:
        allText2.count(i)
        d[i] = allText2.count(i)
    d_top = sorted(d.items(), key = lambda x: x[1], reverse=True )
    # print(d)
    # print(d_top)
    word_count = ""
    for j in d.keys():
        word_count = str(j) + "\t" + str(d[j]) + "\n" + word_count
        # word_count_reverse = word_count[::-1]
    print(word_count)
    return word_count

def print_top(filename):
    from docx import Document

    document = Document('d:\\Learning\\Python\\Phuongltx_hoc_code\\pltxpython/'+filename)
    # print(type(document))
    allText = []
    for para in document.paragraphs:
        para_text = para.text.lower()
        word = para_text.split()
        allText.extend(word)
        allText2 = sorted(allText, reverse=True)    # sort ngược, để khi chạy for key lấy ra value tạo ra string word_count đúng chiều xuôi

    # print(allText)
    # print(sorted(allText))
    d={}
    for i in allText2:
        allText2.count(i)
        d[i] = allText2.count(i)
    d_valuesort = sorted(d.items(), key = lambda x: x[1] )
    # print(d)
    # print(d_valuesort)
    if len(d_valuesort) > 20:
        d_valuetop = d_valuesort[0:20] 
    else:
        d_valuetop =  d_valuesort
    # print(d_valuetop)
    word_top = ""
    for j in d_valuetop:
        word_top = str(j[0]) + "\t" + str(j[1]) + "\n" + word_top
    print(word_top)
    return word_top

filename = "small.docx"
print_words(filename)
print_top(filename)



