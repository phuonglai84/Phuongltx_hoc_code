# Import Flask class và hàm render_template để render file HTML
from docx import Document
from docx.shared import Cm
from flask import Flask, render_template, Response, request, redirect, send_file
from db import get_all_log_post, set_log_post, get_one_log_post, update_post_info
# from resign_leter import resign
from io import StringIO, BytesIO
# Khởi tạo Flask app và kiểm tra xem nó là main script hay imported
app = Flask(__name__)



# Jinja
@app.route("/")
def index():
    posts = get_all_log_post()
    print(posts)
    return render_template("index.html", posts=posts)


@app.route('/', methods=["POST"])
def new_post():
    title = request.form["title"]
    content = request.form["content"]

    set_log_post(title,content)
    

    # movies["3"] = {"id": 103, "title": title, "year": year}
    return redirect("/", code=302)

@app.route("/about")
def about():
    return render_template("about.html")

# path param
# posts/1 posts/2 posts/3
@app.route("/posts/<post_id>")
def detail(post_id):
    post = get_one_log_post(post_id)
    print(post)
    return Response(render_template("post.html", post=post), status=404, mimetype="text/html")


@app.route("/posts/<post_id>",methods=["POST"] )
def editpost(post_id):

    title = request.form["title"]
    content = request.form["content"]
    update_post_info(title, content, post_id)

    return redirect("/", code=302)
    

@app.route("/resignation-letter")
def resign_form():
    return render_template("letter.html")



@app.route("/resignation-letter", methods=["POST"] )
def resign_leter():
    doc = Document()
    doc.add_heading("Resignation-letter", 0)
    name = request.form["name"]
    phone = request.form["phone"]
    reason = request.form["reason"]
    doc.add_paragraph(name)
    doc.add_paragraph(phone)
    doc.add_paragraph(reason)
    doc.save("static/resignation.docx")
    return render_template("letter.html", file_created = True)



if __name__ == "__main__":
    app.run(debug=True)