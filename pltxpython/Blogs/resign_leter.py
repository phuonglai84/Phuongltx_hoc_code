from docx import Document
from docx.shared import Cm

def resign(name, phone, reason): 
    doc = Document()
    doc.add_paragraph("Resignation Letter")
    doc.add_paragraph(name)
    doc.add_paragraph(phone)
    doc.add_paragraph(reason)

    doc.save('static/resignation-letter.docx')
